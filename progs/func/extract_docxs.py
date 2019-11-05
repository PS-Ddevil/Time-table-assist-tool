import pandas as pd
from docx.api import Document
from openpyxl import load_workbook
import xlwt
import openpyxl
import os

def parse_all_docxs_table():
    path = './src/data/basket_elective_docxs/'
    for filename in os.listdir(path):
        # print(filename)
        document = Document(path+filename)
        #read from docs file
        for j in range(len(document.tables)):
            table = document.tables[j]
            data = []
            keys = None
            fname=""
            for i, row in enumerate(table.rows):
                if i!=0:
                    text = (cell.text for cell in row.cells)
                    if i == 1:
                        keys = tuple(text)
                        continue
                    row_data = dict(zip(keys, text))
                    data.append(row_data)
                else:
                # pass
                    text = (cell.text for cell in row.cells)
                    fname=tuple(text)[0]+'.xlsx'
                    fname=fname.replace(" ", "_")
                    fname=fname.replace("/", "_")
                    fname='./src/tmp/baskets/'+fname
                    if str(os.path.isfile(fname))==False:
                        print(os.path.isfile(fname))
                        workbook=xlwt.Workbook(fname)
                        ws = workbook.add_sheet('Tested')
                        workbook.save(fname)
                        print(fname,' is created')
            # print(data)
            # # print(len(data))
            # print(fname)
            workbook=xlwt.Workbook(fname)
            ws = workbook.add_sheet('Tested')
            workbook.save(fname)
            wb = openpyxl.Workbook()
            sheet = wb.active
            sheet.cell(row = 1, column = 1).value='Course code'
            sheet.cell(row = 1, column = 2).value='Course name'
            sheet.cell(row = 1, column = 3).value='L-T-P-C'
            count=2
            for i in range(len(data)):
                sheet.cell(row = count, column = 1).value=data[i]['Course code'][:2]+" "+data[i]['Course code'][2:]
                sheet.cell(row = count, column = 2).value=data[i]['Course name']
                sheet.cell(row = count, column = 3).value=data[i]['L-T-P-C']
                count=count+1
            wb.save(fname)

parse_all_docxs_table()
